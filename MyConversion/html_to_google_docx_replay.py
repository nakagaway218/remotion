from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import html
from PIL import Image


ROOT = Path(__file__).resolve().parent
HTML_PATH = ROOT / "頭皮アートメイク_痛み記事_GoogleDoc再実行.html"
OUT_PATH = ROOT / "頭皮アートメイク_痛み記事_Googleドキュメント提出用_再現.docx"


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="CFE6DF", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_paragraph(doc, text="", style=None, align=None, after=8):
    para = doc.add_paragraph(style=style)
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    if text:
        run = para.add_run(text)
        set_run_font(run, 10.5)
    return para


def add_caption(doc, text):
    para = add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    for run in para.runs:
        set_run_font(run, 9, color=(111, 116, 121))


def add_picture(doc, src, width_cm=13.0):
    path = ROOT / src
    if not path.exists():
        return
    with Image.open(path) as img:
        width_px, height_px = img.size
    width = Cm(width_cm)
    height = Cm(width_cm * height_px / width_px)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(path), width=width, height=height)


def normalize_text(node):
    return " ".join(node.text_content().split())


def qsa(node, selector):
    if selector.startswith("."):
        cls = selector[1:]
        return node.xpath(
            ".//*[contains(concat(' ', normalize-space(@class), ' '), $class_name)]",
            class_name=f" {cls} ",
        )
    return node.xpath(f".//{selector}")


def qs(node, selector):
    matches = qsa(node, selector)
    return matches[0] if matches else None


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Yu Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in (
        ("Title", 20, (47, 47, 51)),
        ("Heading 1", 16, (35, 119, 103)),
        ("Heading 2", 13, (47, 47, 51)),
    ):
        style = styles[style_name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)


def add_note_box(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "EAF6F3")
    set_cell_border(cell)
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run(title)
    set_run_font(r_title, 10.5, bold=True, color=(35, 119, 103))
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(0)
    r_body = p_body.add_run(body)
    set_run_font(r_body, 10.5)
    doc.add_paragraph()


def main():
    tree = html.fromstring(HTML_PATH.read_text(encoding="utf-8"))
    doc = Document()
    configure_document(doc)

    title = normalize_text(qs(tree, "h1"))
    title_para = add_paragraph(doc, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    for index, part in enumerate(title.split("？", 1)):
        if index:
            title_para.add_run("？").add_break()
        run = title_para.add_run(part)
        set_run_font(run, 20, bold=True)

    subtitle = add_paragraph(doc, "Googleドキュメント提出用", align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    for run in subtitle.runs:
        set_run_font(run, 10, italic=True)

    hero_visual = qs(tree, ".hero-visual")
    hero = qs(hero_visual, "img")
    add_picture(doc, hero.get("src"))
    hero_caption = normalize_text(qs(tree, ".hero-visual-caption"))
    add_caption(doc, hero_caption)

    article = qs(tree, "article")
    in_faq = False

    for node in article.iterchildren():
        classes = set((node.get("class") or "").split())
        tag = node.tag.lower()

        if "toc" in classes:
            add_paragraph(doc, "目次", style="Heading 1", after=6)
            for item in qsa(node, "li"):
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(normalize_text(item))
                set_run_font(run, 10.5)
            continue

        if "article-intro" in classes:
            for child in node.iterchildren():
                if child.tag.lower() == "p":
                    add_paragraph(doc, normalize_text(child))
                elif child.tag.lower() == "figure":
                    img = qsa(child, "img")
                    if img:
                        add_picture(doc, img[0].get("src"))
                    cap = qsa(child, "figcaption")
                    if cap:
                        add_caption(doc, normalize_text(cap[0]))
            continue

        if "cta" in classes:
            continue

        if tag == "h2":
            in_faq = node.get("id") == "faq"
            add_paragraph(doc, normalize_text(node), style="Heading 1", after=6)
            continue

        if tag == "h3":
            if in_faq:
                p = add_paragraph(doc, "Q. " + normalize_text(node), after=4)
                for run in p.runs:
                    set_run_font(run, 10.5, bold=True, color=(35, 119, 103))
            else:
                add_paragraph(doc, normalize_text(node), style="Heading 2", after=4)
            continue

        if tag == "p":
            text = normalize_text(node)
            if text:
                if in_faq:
                    p = add_paragraph(doc, after=8)
                    r_label = p.add_run("A. ")
                    set_run_font(r_label, 10.5, bold=True)
                    r_body = p.add_run(text)
                    set_run_font(r_body, 10.5)
                else:
                    add_paragraph(doc, text)
            continue

        if tag == "ul":
            for item in qsa(node, "li"):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(normalize_text(item))
                set_run_font(run, 10.5)
            continue

        if tag == "figure":
            img = qsa(node, "img")
            if img:
                add_picture(doc, img[0].get("src"))
            cap = qsa(node, "figcaption")
            if cap:
                add_caption(doc, normalize_text(cap[0]))
            continue

        if "note" in classes:
            strong = qsa(node, "strong")
            title_text = normalize_text(strong[0]) if strong else "ポイント"
            body_text = normalize_text(node).replace(title_text, "", 1).strip()
            add_note_box(doc, title_text, body_text)
            continue

        if "faq" in classes:
            for item in qsa(node, ".faq-item"):
                q = normalize_text(qs(item, ".faq-q"))
                a = normalize_text(qs(item, ".faq-a"))
                p_q = add_paragraph(doc, "Q. " + q, after=4)
                for run in p_q.runs:
                    set_run_font(run, 10.5, bold=True, color=(35, 119, 103))
                p_a = add_paragraph(doc, after=8)
                r_label = p_a.add_run("A. ")
                set_run_font(r_label, 10.5, bold=True)
                r_body = p_a.add_run(a)
                set_run_font(r_body, 10.5)
            continue

        if "summary" in classes:
            for para in qsa(node, "p"):
                add_paragraph(doc, normalize_text(para))
            continue

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
